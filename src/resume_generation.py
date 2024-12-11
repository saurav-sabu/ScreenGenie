# Importing modules
from langchain_google_genai import GoogleGenerativeAI
from src.prompt import resume_generation_prompt_template
from src.helper import extract_text_from_pdf
from docx import Document
from io import BytesIO

# Function to create a word document
def create_word_doc(content):
    doc = Document()
    doc.add_heading("ATS-Friendly Resume", level=1)
    for section in content.split("\n\n"):
        doc.add_paragraph(section)

    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0) 

    return doc_buffer

# Getting resume analysis
def get_resume_generation(resume_text,job_description):
        
    # Extracting text from the PDF file 
    resume_data = extract_text_from_pdf(resume_text)

    # Creating a model
    model = GoogleGenerativeAI(model="gemini-1.5-flash")

    # Creating a chain
    chain = resume_generation_prompt_template | model 

    # Invoking the chain
    response = chain.invoke({"resume_text": resume_data, "job_description": job_description})

    # Creating a word document
    word_doc = create_word_doc(response)

    return word_doc